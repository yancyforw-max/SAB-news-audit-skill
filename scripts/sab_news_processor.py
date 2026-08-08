"""SAB星系列新闻稿文档处理。

阶段 1：只实现 generate_standard_manual()。
运行时不读取或写入本地文件，DOCX 通过内存二进制流返回。
"""

from __future__ import annotations

import base64
from io import BytesIO
import re
import zlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


# 手册原文经 zlib 压缩后使用 Base85 内嵌，避免运行时读取模板或文本文件。
_MANUAL_TEXT_B85 = "c-pnSOK&9IksiDkp#Pw~0om{Sj>a=f&C1JxHGs7B!oaL-!1w^eyDE}JvdFrS-Mn9#m*%a>CX4K5lf^1t0M9>>n3+{8{RPJpPn?sHSyk+o;SZ2ru(L8xo;)W`T%U;eOJ(5?zyGInV=w#KPv+Otwd>(}KRbKL`g^b4O1eLn%+03t#bmbem&)R=f0fne)1A+0<#cd+_SeU)boDk_nos6;<%wVaYFL@dR?Y^E>16v>ew1<lT-k|3|GBc8kG?fCtkj2x8-J-((|6a&>?Zzby_4bPkL2n=?s@IDlJ;`ao&NQ&UY~c<=2g18l&)>!tulkp2Wflh^?6%H%13=%OXlwe2WJ>8e$jqA+rAofZqtpHT+dD`+09qJYH&J($+j1=?QT;0n(RFct}X|i=hyz7yz}V=AEVFx`d4_RuEJ~gQeG)fCi9!g{OoJD_u8#zH+`AbuYdKI%2IZqtE0dFrLvsv9;dw}&RiZ%?!Tlvbr~@`xEh>P)0?TkJ~oopp4|M}T}WPP!-IZWKOG#%YES8SNxwI^x*8rl$Y5NCtaXhsbVynMY-u)mI!x!6WLmaq@0+hrJL%d(`gK)?m#=v3wPdxD<{ZbCdj@m2GR^ps=NcbP+8>kW8N907OqLpaNpiW7?3^XDU$XP6te~v#;N==OTs7qV>HJ#u*h{XLliL03)BAMwI;lL$Ytze*$zz)@OzO?yP9<BpNv=M=KJI6-ze}^j>9gVb85ZK|NGFos)nzZzm6!BOKj}OsOUv@M;o+?OJ-GTZ=<F(K$fw@Fz8s|04tBHO9xNTpJ7rxkwcEYn?qM=pORgT0ca5a>IeDy*h)9*PMz6g^Sxqc^Z)I5hmOP&ekq$q$hMS*WA2*WPyi($Px;vf!Q$K}k!;3=8G88vX=AQ1H%S`wwGBVcYd3`W<E}x#Z&eF9@S){c7DLFb#7cPc7UpN_l^w~kOyZ+N}a(Z38m1Rt>wqBo_as)`e$^1JZHX+Jxr9j*6-f`6T>tB@*;6_JIC3iudl<TDDtS859;k7Ho_nvRW4(@l8m+7qjkWD?vgG%9Cd3kzRd56tfzx(Us_OE|sOD^Z3cQL%!NLS~1SysHiI-F~Q*7cNP3v1ov!=kL$@VtS4&TqeVx6+k<vMcPQCflR;%nx3+UwgewR_&cEXkU26;B<C)xQrdWS(O>2?=E!(!7BQ9$<sSIg|erw{bPKz&YgG3*T^Tz*Xv>J%UpPnxKo4ar=)$J?JRN}Sr7TvWa%WC>&Z!becY1cCo|$)u=oe_ug{CJp`1>3S(gI=s;f<BxBEJyc%Wswvd)}MP+oFyg}v%lUc2X5Yh7r0BB|ZI_Sb9;b(o~xk?)XGEK?pl*OG-#f%!FFpJqAhVP#)Vz0rq~Kze`w`f#IDc<nVvD9OTETB(jF4n8IO%AHIXj)&Fvez#=l<)5<s@z4khkcF{(z(A^aaCKh(#y<5HldBJbu}~cFvy^z!N*_nE-PQFxm2JmX-yh2#$@Znv!Yprqe9(7r{gs$F6Wso^BfKMf`Vub7eKI?6`GqOD#a&Ks`?3J|jE%eW>YL0AOZH)hx9n8pmgMDPSlJ^*=dac8<z~(w;Q*^3l&=U>0y0XrKV+9HLT$ogvX#Bp-t23?^7=3z$4IAZSfQPI^8Wm_w=FVexN?=Vr=?cbzoZxt>g2w4>)baUh41Ve*y78zouF`b{87@gQ<2l1951G8>vHIl)@<CyWczx!^$aFksi%vdl(#ozz3|7==X7C19?!l!XU892`}e}t$M0&geK-_4bCcleo}VQTUu3z5t1s?y+-~lepRC_yW%DCH3a=loRQN^iGB{c5S{NG+sm@pU`QS?fjBRI6C{k`5y!`lfa+%_Grj6xv9-2yWGWGhj8rO(oH@{6OtF&ZcUz6=^JCDjEzq7CD;_Tq`BN-+64IXYUWmAvxmSkZsa!)x@S^qsIT<gjt)13|oOV3EI0C2P0o#dz~`VnR>e^>T~^;`ZuSb9o!uE+1L$U9M)WDN)R-#8Uraz70jHrw9dPW0+Yb2VK!#DDX<lP~M?VRpAJK%|8*PnovqE#Ioy86T4e()xpJ6y~e4976Z#W^j8vxID>@dSPvTa9_dyxmBV+twOR7=jM|3TE2xk6_r2Ak(;YBZSFePD<0F{iXGGaxpa5-d+b%^t;n?D%A&}7k^Yn>B1E{n!|hq2VHLT{*~Ox&Xj`h@h)Oqy9pxFTb>-PN_o=p(EORR;Lb=3)N{!p2Cr4z5=(2X<Qp9y~GetEE#M{11E|;*q%Nx`UUi%+K!pWy`(l+=#xZ0kGtA6^llPo|F$6&kfl80M4_rjog<I;?XWxgisVTiQev{!Fs{ZFJ#eCkn8z71mJv4>sO50(WY8{rdft-g1nVHP+QAn7~|ZVn*WE03a@C6_0!5Boahb4*EI6Lx?kLzO|aUTDs`C+Ugk6!&~#?UI{M4d6SQFLIps<=Tf%Y4G&KHx8;>+02QkS*jaA(UFV8%GTgzPL^4woix_$oSx3g-$JC&Mh`#8N^vO2-w!*EF2XS2&R6xMW_$5o6y!5CD4uzU$zSJXXU^q>lf^DfrR&0?-zPU8WKM(UW!rEZ2pJvb-qTGo-G&+P{H$x(S3_=eEBM{A>iljH)%z-_d~GhdT(^t$mvpfvkI)!|uHV7(MuY9O`(;pF!5L6JepqYb@r6C1Vw0HD*bS;djh?U<5$7_I5$h?)7MTw;ekbRX=2y{TD8%K2|G6T}VkPQMTvzv~i8@i~b^SmnY6($n_yK}UE)78DX?!oRWib?zhnK?Zl~ay9U7}e}(S>&_XBwTnZeFBC?jmW?Kh3Q(a=%c+O9pr^8fJPwtA|Ui>gR`SwK4tx#Px#_lj}23YlBWrbq%{0M5if`LedMzoFfdir+0BMAx0S?DlU}8kR7@ea%4nylpb`OptIt?ZC1v<F{{9mxvcm`J@S;}<uUa+n~ahJzXG~Wpo_nGGC?(v?TX4A2Z>5WI<uGVy%%;{8c=3}r-Jx`QTYK+k+j4r5~h%^WKLHj+vMip_X1m^y9SYRvpTvdc)JQ?Dg)RNoeY(hiLVZvu|@tus?zN!@l5g#oKhW!zhc<grf_LH{j)U@JJ_19XLLq5)+B>Ir~{!T+Qb(KsoqK2n8V%68_Cn^2oWpmu8HvE6Eu7WSJ!f{(l>oZ^cJPt&#=ZfS~PR=UD<*w;Cfl4$Ulf7NTwTB!T<V`|2Ys>{`1d2%g_25KmUiHf04<XW2Su~uZ{22$G{+*Dxglchh$0CYU`(;MkesnKVS`)W<&-E@5Ebh61odms3j4?mvZ%|cr{Kid>x)P=Ati28ntxsbNcN*Ia-UyVywn51Tw13lb`FW=+dDhCELDDTV@O=!tm-%zkT%6!HT~6Fx;%kgTuKtUQ6Xo#6hiz^V8oOTpiGdLoWe_l9g27j)>cA4;)7X7nxoYuC9__Yy$g8c14No(0k>r9xRTUm}d8r!6;J^B?U*s8Ao)Al91_R5_O3lqCoC+vvI=kLC(MtN()?u!E`UzlLu8_udDD#!1#L)`F3?6$`Hjezs{{9c3<C$*X!bOU1jP*Utq8*cb&89=ifweEoTQ(wA&gUe9J^3<O!5BFWOw^l50mVYxG^|qkN<Mon1bU7a#NwKr`Jv#LR*d%bY2_6+QMC3mzNUhmBTRuZYqT-B+Tw2dAu)<;&qhAh~!U@TU9@lc4cJ5Q>m);ni`|_%U?hf;uigik?CplxjM?uHHMNWAI0M0VYUN-X~044+1qKVQs?T)RD3X4kS{C`fkqh$gx3Z>IJbWD{b-6s4~%4Ep(B?z^~bVpMIDz)vC!G(=P|2?o%L$;wdx&`kOmX=a&ba8AwP`NRQXmMa$=sqi%9^%0cB*(*4IM=4Gl-Lf$vyJNn1jr(XJX8>g%-0{3IO{bJr^bwlyR<2>R*%H6I^kdxN!mXE@2Q6enKOP0wA;HOq|N<KDe!<hW{bS3`#c$o`^rKH<vshOQ#%aR}fu^7n7#&;C;Z|ncIPPY}u)qktq#Q}AM36BNd$f@8Gj;!pq#v=G`x`cH1GGQjN9j)`h^M?s`M+_#+?+-VRC>&JZeC<9b3smxl2cjo`au;|9r|<4?7m}UXw12?+oJBwp=}7@|9Y-$7QbR03>YHRjCsV>>=(_=Bz-J7nJW0;>Unz!E4rK&;*d{w8yNYhy5=Vgbd(bely4ngZNxETh66ZgYw#dOI@E~^|`E&AoC5iw!1)l`xR|SFxKKuMx#f6;GU4D+(DX+aFA!=@g@apwjThZ&_tmHXyhgROk;RkbH5n7Q|(J}aV)W)2fa`a%drv{a2nTeduK)8QOcjo^j=AFz;NsSC*=S)e__PRl*JxLL~guYhYgQN6Y6??G1GPpXkpqUc7Oo6l-{33pdCff&kldRxyZVpx}B>!UZD4Wal5?aomL}1DK%0~zV&cpK+{(%GXr$0sC@}K_nC;nIG1VR;p5c+7;lBEk;#Z(}6L<(Hf)pRSm-GPkht|7X3rusZJTV*3OlBb`NBRMY3bYtIsR4xt*(FqSaP3J1GqoN1<ttBbkxBnot%~Qa^87JlQj01oyUUt()Rlg~_s$_eYB(1B%!RZt~LiJWg9L?|2Oux<ycfQcaitW;41h%Z=Tl5t8e)Kte!lW4wp_gN7do%0`C0=y%gO9RV>D9OKU&veOjK|qfb(#mGAiTkn+(uButsm=VEAR1(Jcjf;aE~nPn`(@3PCQp`Yu%0-Pg%I5IfF6ssG;e-7{;_b7nB_Nu^5w_(%_gZE0@J$OTwwB#)X(u(s_W+e6uQ=6ZuBks0(9^`j@C)>Fx`|sihEOQagjzwep<Qj!nH1npHaS27FOB8^1q3=wkd)hSJY638$mjCr;Or*>7;N>sN5v_%vaI7yFQCN4M#HKkNfevpx;q)>TKhq9Fn;y*%XC51!-=%X((xG@+(0wKUkgldW#TCw8<YY~5Bt?8Jqstap;t=S5E>P14L$4LDNn@I=<OJ%P)+uLV)?N$Ne9ALZD<$3=Fl{P3Uu`Cqb?et2CfI$86O55YW@6<BJdm&bOU9PYZEY*Xmv(-D&~$U39V)^)Pelg&&wJK=VYtKWkmpvx5pkTWCpp}bO~isn2daST^B#q@d8nqi2ZGBap|)KaKy|L=RRu*l8~UN(e+gF#k!!rz9CbtnbscsbMH!B)Ek2U_FwcARlKl$eDY(v4@nO;ikHO!`9j4~c!)IVW3la+6eVhfGy*VVObD25U^$z?N4`sB<{}TG!d)=j>upUC%wy7~motXVb{nL8kWJJA>$#-*{9>B4*nB9j&cs_2{l8d%f(SBD%1w`UIiK`^1hDJy^95&Q>hbpZ;F1c^c|(46cs2CO;>Yrkseuy#<?8)iMj@>f2!H37@5d!tT}^gQ1ddi@ROO3)+*uI(z&yc-iJH`XfT$2MxJ}WX3qToXXdyP1rzF56;A~FM(TVZNlkOp>Hu!je1k^@kGPn+&dXuPp7N56ms>4Wd1_TUB&%?&yfG>%xI1!M{od<iC-Ia^CrDN-I4PMzcsB5HJs+u=RD;ez)8~^c^7Af*@#XmjP?lXBy4wCF_D}4M(8m*U0E?1Cx+^A6g5=OQz|B{SoS0W7qA}@Vz2!^!cNt8cGB0qk2EB71@p@mr7;v`Bq##hdG0Fs@xz_)<E%dq@wR?PfJ$0lOm}+OMF$auPCux|W@hL@2Dr&s64g6|5##gMjq7vIWEEc@>b6RYvoeH@T#<($<WvDutFxl|WJ8tlI;_siI!LH{jsrUNkPj?Ua0<W@8G$>bcj1`ruH*M^Ej`|%-ycr|T~dfe<^jvYWCf^Uem1Fp$$BqiXu}Th!7Q9nDr(JZXy*Czw7l4674}&KTlQr=GB1z}5i|YwgR5S&gyo1Fox+Q8+Tjz}v#J=lVkF-rtv(%S?sh|-IZ;RY;gLvppVQTk&YBTH0Nw1VnLS-ikm~X(xk>1D=M17&mdzqC3+b8#(YBwn`kb6ZIP%;rIA)@%)tg`$?hug1!=s@7-b!|RI-Kfrr1*=BQu_*NLTsCER|MV+u09C`4l1>DeIKz94Igm;P9k}K4y8$_;sUQ&7^MIRc-<|o*#wH+nFkX&sY1XDohg4e2SE_QhMdcHo~n#k)4WF34YRs^l3Z?xwqPH()DoR`@VxBy&)hU8v)T4q93Y4bjCKD73e?rJY*=K>3}#@JMIHvyB$se-B4s99?R2J{t_Zu_@bO6XSwIU{Wr{l4gax5{Lv%WKD|VXlD{ix}vg$Xv*{lu9)9O|-e@(?&(Pvya9-^yCdQ0^=du{~_#2B9KJS`9qw`|G6P4fOMS|$BsY{Yg~P9gSK^a+j84{mR>UKfJvV{LeR?UvI*f5L-zE~wGdYMJoWj{V!WIG??e82?<JNuF*fZFFlbjtT>Te=f~XnFoy=TtLH6VjbLm1k>)}=qJsuLRhAkA?S@Lr;shYv1&16O&|8i;pn;M;|yOoz1>qU-MQSR3F35wnn|!lI*qOi%w2f(&OF_%9%7Y>ws+rwcvROupL1YDg-?Qu$wNt-z&F{v!OL<yT8I@t+_@*=RqD0niz9$czD0JyyDQCCYPdX1ustO|(JBRxHGD`zv_rgOcGFNE`XJ0vy0T^BGi?z^-Y6(5z)<o=R2_s&falx<XomK;@WIa~s<g^jP8h>7P~C>Yn0_zwT<uh#v)kF8A`;VWC1l>kzF058hpg_sqn?Zx73-wAKu1M<1j;X6K+9*MI~Nt}0x)>z98__&W!&^ibc7Qu!E{?uia0g!qaU3r8?W*<8Q3@wd3rv3D!PC78`Dl@(Vuoa@kK>f(YwNh2J^Eb1PSAvgpg2F3^hMWhn;4!a3J;!J+g*wH6*Zt8w@fY!I<V_6sk~{C?YX(q5{I&!#_EI5^UP-q|aZI_BB0NqV>b>de_KKuLsxf#cUdMI_`EH?(HXpmWWJ#I!qQ2ZAm|$3YViDGu->6lI1D6f-7Du4pB)oV%0>frSdGFfUkGn&)Qd6VW2?2Rg?2~3ggxF2*76yHe2H|4nMYdQK(8pk4z-HT+?uQ4=d2#P(a$0s>kWo0{taemDM&P#!4{^wkuvk9n5^2uKVEnEUmT7I0KTI_~&Rz5_@G3q0R$G<x3#HnqQOJ?VvLgl^&}8r|**`C`E&6Lw0onPp((gQ{2O?qybT(I>3Ad!!DV_p`xkEp@d5&%TEgBCLpPiyHTKuf`Rf*9e{tDy~9KTuA+j#aQ1X;j>==<o+H{Fj79|?IL3uWzSdyE$~^T{ZplRiY=hST;}Kt>gP8eoo!_D};43L46)-~_2|u1M4x#cd`6BQs4$MVR$d^kanBBxQ6<!ZwG%wd8cO*xeIuBB#Cc+i>s@u%qdc9LIwYUfEb1*hBC5TcoZ^)$vHa{3ZY;TFri8(zu6IE#?LgJ}6r}d^<9Ee^eH8D(j>1GFjFMVpzyLjzCIZFheusxe>A8JfW<=KkJg&z25z!9sjXsn+x5QGB1^MT?Ty)hz>K0SQ8$h0W@U9y;-ApP2E7RG1<El%R!)@2kkcw{w1$UUwOzTFFBK_Z8$$}q4=d!5eU^dm;Rst9v3zOj|AbYydtMI0ulD>9Pd1T|hI($(@6^q;Q}CqfTA%3}7Vhu6vRx}T7&f&cojh`D*8Zeqn>Ehek8mH;5u=EBi;B2e<i;3-0k#|F|_g~ds?ZsCchTTpcZ@R!!Hh0m0<<GoCsh_3h6nFwg|A{#oa9ztqz;;-t^l;nh1wg~v+#RouC<sJ3Ke4bI>>KAfz5LK?6QY3^KKsQ=y0>CiWSBTnT+Ce-d+4wRL0jp1Q<v5>qb9ge(Fc(><xk#@N`6rLXx;9?G{f=gY0PIZwcn!CfvQy2ssjlF&i28%GOrfbbNn&AN@Z*9$i;$nQlN6MGnl4m`3#QxlO&wWXz1}BdHm{DKt=v##a)HHE<}pbZ7UUu|X7|Q`;f>yK#!hYz!NXm6L4}yRF!(kDb<_tS^mUhO+3InKZx~H>r@%)jkn(LVSxlrHFWU5!?n2NQsvSm;M#iK2K&HO=497!fg&(r5f=Og;V~x|iTNFjWr25OXA)3F~;gkJdUWhSDzT^H2FY_AhEIIykYDNpoB%t02^KorXk-&hjFdPBFP9?C>(_o^BX_|z1WTx#5L!T5ie%E<NdN;3;BK4h$+Fd*Letmj)L~ao;EW$yag2|##MAbW|NeznDfchoV7@SreVbk%iz+xwRX1_vB)$_rL_~Z+&k<KRiga|(7JKzXYD7eP4Zo$lCHmjS1sV`3NqahLH5w%9+iLJBr+ahhR$Z5J?(mlmi$eC}nlH)o4Z4eK@aU-e{a}#RYhE83YNuJ_i83QH#3Ul`fYqH2&Q`zHt2e={Fogqfc<BMTa;HcGx<k_#yldsqsk}ZhQT7Pi1GSR(Y1*#~*#jv|Hffa0CVZZbNddD;;6Z&j6a|bMyYzb-o2_#Hnaat%J1)%a|E&Y5zghR}*p%21n>+2aj8k$H{P-H|YJ^~~Y$&1h);=0(h&R&8e3)&++X3#X$M2vZ3dGuU;5K;TtpUw6SB9I;LW+%#B)!iF!p>I3LX-q&g`@KDE@%3`L-Znk`>W)SPV&g~8q8_EwVrdE&uOqG_+nnqDUPHG?uclpL8QM!YT<U1x5%X>UM&opi;N;osr8}QV$dlY1n!%;V1rF+vriBG|4=2}8!?`w{RU!g^_s4(t(|@8$`|Dr*@n8QtP~6va>HP6u|LcEK1$H%L{mGU8@n3mK*2OOUm;aAfLUPU@|Mh<<lu^u~RlV?^{`>!gs>`ejB?|0t>nubWl9vmb|LL|+)JJ00aBlSMn#K!YPc<m{h+wOOu*#4uSzxq(Ol2`qa45St-s5g5vbs<jUSR&r7JsQX2d#S~ix3BkSh8g-9nQ$%OIPkhpBhXzM5z@426X1T#y|h_zfkqp=ZCf3VeO>^byu-W!~<l?^MeSW!4>RF`Fn6Wm2~#(Z;O>f%7}Ql{}R`O4owU6^Q35I64SCd0X=Pa(e9q`B<A}K4vv#rjbTRzu`*wY0!LuTcP>SppP=d;ZQ<naUCP~{tLtU?=Wz8!{^`J7$x$=io#KG$#<Q?U^@Sj^mfn4f8n+tTj1|z{V4~#UuI5JMyI%Xv81i7yftHh|vh@eg%N7=R*U%L1CixU!vmp}tkVZb7&Fy2$qr5x~&LWz_=1F^6&iXrwG^h@)mxg;^`8V)wkNf8JSA|ekMpU+Bt)F{)>dC?<Vc8HKFVjliRi1IGnDkGXfoTp9LrsZ_pT0Br{T0kxY#KZp>xkHVtfYI#c2<=4E6b2?5`B>8AOzy7*?J8NkI=+pSOO*lfis*gX8rlFjOvGnt##{S#vbXwWb+A>peHF=`b=_ydrs-q6hvi*F`LvdsH<4br_W^Jl<tHhMtNvKEvM16J5+MPr;l#4lbR-RiT2!l?Y(0vfY7Y*eYFP^jChJ;Vh%!+@r^oXKvfDcCt61uYKp-pK18B&Q7Ey;<y%2zHf`lCje1r)<?2$0m*o-z4ywScnN-cMN%zvY)XkKL#q85d+I|-!OYpn{$bZencBZ7f9qNYcBx9Wqi_icA(gou01la;<_aFu{I}h9qI0HMkmZXAgl@Ej4GvO_aiqk+LGw`auX>)mg{_x$r3~65`{IUKp=**~sd6um$8`@3J4tBM9C_F#izRk|QCMU<LaKlLTqVN=sJW;2T?GMR|`uwlghAT@PII?egmTZZ2k7Sh<g|7s$Y<?+XoQc2?sj?(Cs?7Wlj!UC^5)`60I~AdOma+iByj;l^@t?f0-xGQ6NuL}&J@_c9820=5MH9oIsHf|PIOznIBf5_7dwtkXr#n6-v~tEg0**(FtJSd}5=hrM)L-EFqf6DkYzL@<w@^F5=V6}n*Jqja8mUn<Z5f3-Bs(j5he7j&;9d@tB=8Q}+hi}4d#)CPLzDTOvZ9yI1Y#K-MSDlID_6ZCAiK<eYsw#Fpb8xF;}IxUTpIcq%MN`Mm+0^LLLTY%bF}Xvz@OJWl9dh{3;7?f5AHj(#70Ys_?jKqQuVQM5U?NzsHH$7Lj<&0A-miyCK$xx7*<iDuf;L)a~ADHny6eog%)ZsLJdZ0_UZ0{e4;$<fH3yih(mnUP&-gH<O9M4WEg`C5Nn|^VGH#_y}!&WTZq`u!XFf{*-llgZ`H)C5{j_9(ea@%gEdqr-@!Y>yO260CPR1^UZm`CGOvoDO6(o@(-|C8N(b?qk@2ZaQVntNFvNt{uE7=^pHui}5ArdXqr_L(of>!sE7Se;m`DcmaddX)z!&>%6KPPO)%2wwCmbNZkKPO93OZN?3E>_h=v`R{OL5k1&@<`%WmP&zRlaRS2a0LY>OmsF$!fv4>OC%g^Sj?L&XX=S)8!c~6eGS`Zzn@7ghmMD@J6Q3%hJB(X0;F#<%T|B$XKBQwQX2+g;@%8`;J9BwEB<|wNnf>`%ZR}X)30>>h%C`!jYGO81kl~20a+_FcA66?x0hpY>xpW!aJcAQs*_hAO_H^<cE5{2yS!$m##2HgL&THt1Uj8uAj=Z)ZvI~Ihc_Mbe$*z#qDNt`4NGa-OAu&&7iBm_d!n(HHEXPKWGqB!+<+YIN~?N%zx&fmPKuGWRyeMIpZ&Uk<p5<$PbGn&xx&q&{lUkd%DunC5`M-(<aqS5$i<f-dCi-+@LcX<B3+i<vG3}xMHgjf$mc?48EaJaEkVWE5-J|<Z_QT(e`t`q;wK#xCRG2+`uqr@^F_!{Y!;)Ivus0ps>1tlf{q2`pVRQ2=8gIst5!bgTgFeE{k4cx})h{;A)fF7^^m0@pMLxpEB@#ZX8(8p<&$JkM!+8Skqsnp(N9D6g-m5Uk{&_%st?qlU2>7w-20~TTj`lBMl?y)!TsE@UP;JgTsqq^(t)~81g+(4xdzjEQS0B9KezHq}#hmtsd`Ax?Orx6+z3(o@a3R&G&W6DgaRT;P;LP2VmfI@ljqEp;SB@;368%_6MBAIe0#cd^v!`lIE%yar&O;oaObtAy})p1TEXN2qa2>8UH1qLP6RtQe}$r8zg62S34fA?hH=f0Xrr;RTTs2JxQGrl0m(b93!yli<peCN*m1aD)P0+JazMj>zCQiBe$UbK6yW%e%`lYDJ`*_2H&pb4ct;uZiLxg-zL|exi|WD4hrDgJvW%z0C^0c4dA>iq$4CEs$O>a&Uj3{0Y=Yc1|}19VbwS?BH?o>Y{a7*T@3~VkSN(%0Tx1^lpnJI{coSk!$SXJ5f5gTlSRnj*XL9DXTY*i6{V}%7Hj9s9}{6|7T9K7GmOLy3dHBP)f`ft75`Os<_reswO+Dx@RV+g;2s7=vsOnD9|q6#yJ`v0>&T4oGSCL}#{z{V{Nx5NAF}?XQEcR{>Ql7<rowJ;EC{OAe?n(J75%GDS7RRRXH^?f>y$oO-~s+ywWeKQFtP$L>Am9#UP#k4w5W#8HqMZ|EsheCdt-S3`zLLAVF3D1E+M1#kTM6~BJkg<6FMpY0q&SYb<!P^>{Bn>e3DU&D}5765E7O#^#ZW|3Y*GS<2Cag<=9cXkf*0;Doj3EVL({Ctr<?%u&8rlA%_wHD*CS$v&Wvl#p!>R5mO3al_G0C_U3G?30XZxaEzbkt`X5F^DWW4`FYqg#k6@<iVQ1<vT4qPZS(w^DEx?gETPV#4{vH{vZ(vC$j&FF0_V4y*`jwxtb2ck7i1xs&dw=Bgot^S6zc>-fbMUf#-xP5z0YBAjOT7k5`|pS%MwI3tcZx-d3cdL-o)V0?gt*x0;^y&OC2@}?<xZ`EY!bu88Ga~a4{?sUy`FFzz4(txtS9kI|n8QsKU>GuB_PqPxODg#CK|m&j(Ej#D1oowiX*Yf*gFfGfo2gVaBQ2@-Nfz3khe*bua03ox%t&_5H5D@UaR)*B83W@SAO31750lLEfsoNJL(`|L8I_Mq%;t&8ycQkImo@Jwr@_Vh%~vv+~%kmN{eiFjh6fnyWGbvs~o8;?}{86ljj`wzYs0W#NqI2@;Cw14}KdCMk1$+@}I&GAw9NTnLEA#ro@G7o?-Y&?gCQKcn)990jLqOmMc^t&ZR6+)%4y<rycBQ&k@09AE`4Mm<61^`(x}I@Ri6Z>U~aN#d8>NHkvajHKf777rLF#?|04GgTA%0<_?q%r&_k(rWrF)*79V?x(a_faqfy^V#LL>^KwywQ@-+T0I8K$65!LY)+8>?U|O2$W2VWQGSfZWjW)@@^-6sPRQ=%Tljc;6Xua9nYC#jgJo5u`SXbk&)>niI}~H_+&~?<`>N$(mbeeu*M-QW6m*1Nsgfgxd~lFm)`#m?%nuty=rGh~p2*xjDJUr26GYufVy#ntH$_{lJCTYK*}R6=M8zd*S!%i5{Da2tAuK`Rt!@H?in*eV!dU`?W<%IyRz9Ove6b<RgD*Q>4Ao&Eo8?8_Qy%<L$>q*I*pi;yxXAeeD$p{jq?Dem{&4fiB)j`q4y1tq|1tE$m_RfnIl(bhFF-Rz4-&E-LuHeg!b#_iT2y{^amo|fMfA^j=k9pRtYF72&5P6yh9OYRygSG8dpi!h|M1nT@PYjiHo#*aQgK9xX017UvVS1-Yz4kFNzV)y%T%&~)Yj(9DT1Aw!X```B{12fwTMKO_J%06BLr>9!7WWR#Hxb(f8LiroQkMlt|1vcW5r)=1JrU-gh*d&4mBaMi$gf_p)FJVo%7_26jiFHop#aYIbu^HFg|@AeA`ZHi={c#D}(Fz@I_ztUwa2CM&Ci)yHm`7u1gbwf~pAbxAKLc(0wOm=#9J|{l$>1gnVj*N~`M81DK<@^8y{pK;0M?^>O*?(I0AA)e09hA4f1cNAY0h$!f)<CPTVwc?P2>_3?`;Bmxn$Wd2cZ^o(jEA*0;RSG!{YKs7cj4y_s&wIvmz^c004X`afiZ~}aQ0w1{!6M`wK@j+_OuX)Wl)bT(eHS<tZS*%Y5QY(73+=`ebh`_Q4#+D^+tJo?jR#bb4GzjKo3=+czL=?e1YUVK{+kT@$Tq&uV+OeUWirLt~3unpL#V6=g=aQP}Py@!f?f~XwZia>g06Gq~pbSHsOTLz@hdQK#Ga1w9_Fj9l$oHw$EEA8)ruq<aR7c1vWwsK8p9W>+C8HvB=|bmJplb$wN`=j7)Mfoy(=SF$I{zi3?~!vW>^GP)Ia^h0X1GxkHIM^15gp@}@v<oiJD)>_7)sn(+OQs>^kgbs{9Hubjw#NdK9{HWYV%u8sNoovmZ^}EfMKy{Av?@jAzfnwGq2qhbBZed&4(ysP?+%ZL1lWdxom|Jblw50;gkL&d1=M)ma<k7Q4^|<h<qg?PeDMIt6E6kyREwNIYSwKm_e-2O6V7&yT+O-+40%1A}Z$xIOCoUCx~+A&OiXGTVb5SD|I%C+T!358r`<yGl*BvujE<KxG;rkus-zS0~AW<%iI|rwitTwW}bTB%mqb1+&F)mV{>#`Vo-HYmNZ%r{zM<7cZkK`p{S2I>^xfystK(zJ+)CsiEa?{?|_S5c+paf=Maxin@Prnh4{SAGv{@H;HV%Ui*5<Z{^I(Lnsziu42}%F-3sSCwd4SLC7ZXHcM!Ub<}2;nAF9gWdsH!U@5%#6EP&~}wA@i1b{?Q{2f@3qC>7O_;RSfkzT6@*m)40~Md8-I6YWP11Gk;NA$y~)=ksp*`2bA{4nJ5(n`wC>kbCWr26P@`$a$eYICeH5$s@QBX@EODt9;=csF8O>Ydl!Twq<H~v+07UrdF936shG}iv3CA6g(QF5N+R@0vw6Y#@JQfDx_dDDlalYq6Jnsd$I@u5+bk!hKA-a#_}4YWipC6gw$jI?~)0V-8m~`$mpMT3n`U33`SZ0o>&DX8k(Wx;c)Nk>jN_5nM*O4T9Zl9MQS}-SmHyS+Ab3Q_rC#}i`5pj)-h8W(}UUc!(EJlhth~i?FJPhFV~Rzs*5;*J?kI<VAkd%JRxK{4z4v<guOSk+<2^CAL^1@EnmIG_th|$iarw@D?H*1z;_$z&NfA7z*JP#hGY#6lm=FBr4~EaVH_mg!pZ@F=qu*K^3C7ZC@4ZY$#ic?9v4&nwR;QfRi~fUpViIvn9MD^s-9LQE4YWOw7oWvur@N<K!R1)S#sC6=NRtpCsnP`46Mga$X8q9IDA8hM_YMFCJ6-8Ghu?3a~vuG#_Ezx6m?AwY)x8w>^fu8BGk8oaFq37<wG&QiRWNRR)>Nhz=|j_*QvzHDHW@4MI8@B>=69s;_NqLV<5}KHPBWado-B25xEqwE2G^%@|to+-hX|Z*PQ)2qQF9}bF;#yCj@X-?vnXyYYn1RQlq~!xqeDsE=Ez&+ztgI^W&{Ry@u!WaE3jiQkK#sKnr|jyly{jz*4j<$zIaE47$XE!<;eJ?aFmo$SF6Zh~l>rU19*fRl>_xx$<UYTW(o493!v?DA|y%yhQVpwi*A)bz#1g;=mhbsOSNOLd*)PT&JZ=fDRSNgbiiO7xs1#Ijp4RJ*ijl$r{2>w~%izV568#^qlpN8GZDr^{gQkWlR9U#LBmzjb1}&gC@;eB}q#~vLe0tf~c1Lui5P*<x5%d+u2zWs^<$*mhg2rRwW4EW!+UDolpsJv_$B4N_hU?6z`<M{KhjZv`iAhzgbFAUZv0IxvXRD-zuEw3PXOcDx_7gR2u;iP*x{fzvFW}XR&TgpvfOy!~fr_Ny>60+uK-<0Kl18SdhP2e}%l<rWq%p7KCFnBPBS??5V_S%$1dMoDye>-Yv(7CRQhmVY)a!Cac97v=*W?4SCVe1^0wkQ?qsIs2<*xxMDnIZ<ZSCtxy^E3_KLX%?nte@Z|iG#n_l#)m#mSHdgAu2VINIn@Scm9S(m3_Mm|ZxV~ef74qfS97l7Xv`+k1aB$H+igi6`kzJjM$&b!Bvl>pfcm^Hmv9(NW{SgrEB`;bWIW*wQcR%#e11_O6J1&3E4<sCx8yOP|V=)02PzMp$E@eyP@Fdq0B0l^L>!2=_(sW{<92S8<0f2=A&Yr%#;vCJ9CPOaG+jO*wh#@K3$7=ein`-4Q-Q6o}^!)a)vL7I>=a5cB^sZf(QsYDs40W4{&>hWpk?oEtFY-71DZaLIUY=z%9Hi^Ps4d&~I2|Gc#Eq8hG%mtlae16nx7=lfI7RJlooTLN-d>8uH}t-EZo1P{h|H<$7~Z%EYMo!r6xj;y#vE=+z7jEsQG;ZVfM=H)p$yY&z>kgcWm_;il?hv~Cp`&~S3)t+6Hq!EainfNA4v!B)SAN)b3KvEjN5fdNlBjCSnCh5WKE_owJ6j2=v0-N+B>RW+K&joO{#f&EI*f;Z^W!cGK8pLgt{lbR0{XGsBjA}Am&(X7(*sZv#(Imd@}SV%M&P3@@qA3JGK{#nZ~T$5B#t;EAsZjWXnk2i$T>gat0CErS$?-PM(LZH6=65>W)*^J+Gj}p_a>F73*}R6B{Mr*V87kdqqW#ND=cke7hHtnwZU3y027V*h*36cwv_UVs2NJlDy3R3$j_P3X=aa3iSx(8DU9vTlRozuA&x@k%dJ=^FN;95T3FiNbvB)mN0|cgRBv%gSXGKCt+^u0YFn7G(Qf`>k2D7MC{*c2}cc#tsDK4ymsFYu3pCFjkeAP%_XNBA*NfY$zcX_Fm)H1zL8N_dnr-8nERU_uB3LOu2<{dWKX+O#7MG&GW5<7n^oZkbUWzA<gGBmR)khysH|56!Z4|UW^cF)1;=(OsA8@J$rox)GuBzN?ej)_I%p2*aCWjeOL5qd*S2yKkCBrDksl!=Z2r2q7!5AQ;AQi*XRSfgdL^C?bG|D5Yf>$A7`Hr8nZ#<K<V4Q4eq>L7`YqH&c=VB@(V0Vm&Q>HneE8>s*cmYPQ)Hu}h;GR~e1@VvKPL;pb}lmNFTV|YH~4lhgpH9!1#8MZqHe~aP*7}s*UpBbAyLtDYlfDVj`nvl((it&QNZEq%i!ru*q<I6E3LvpJ6ftt)ieANix-H(gyaaVZ-LVztZ`vGJMCn*G=BeH2L6w~4FOAjm<|p%|JHNzvSIb9roZA`K29_KE`wm^nt{nxVw)@-^B-iqfBY>V$PLym43Yu|o#e~}myn-R?!d@jYm2BHW<8-$q3u7&m_O-`v2UKXo8U?WIZWDAp}6__c#Fiw$Vc>cLp8Qhv=i~bew!IPSaOxMA|ZoE2=ByJ{`5;+ePm8619BgOf)}WFk0pj}&H{`&Mdjzn>&cx{LHlzw{1bNgC*9(iy>x0XY0f3h7D-XX2+s`#h3XhpiUq6a#2!-N4SrzkZf-ni1D|y9JsRyAg7SUv1xzqz3}UOXV;UsnLFh92q%htVhh>fYK~)<C9G-w&5TrUtUJQdv&Fb?3A0?5KI4LIBU5L$o%&=`CMR=z!V-0RJ(c=DUaKCFvlihS@FI(AzX^zaFPI5R+T5+{y5Y4(-O%<2Hnlfx0oxQ&s9-1ntg&W{d<Xxhr>|7@|71V%{6+MTo!#2`<@2+TbRmMWYrwu>#v93FNNH|7`LJ*23klzhj<dAVEp@PkCfC+X_K(tCF#arxoy#*rAGy%3AB3?M%J=X8q4kjxM0417nT}kQ494t*I?X?kA&JrN?NAjErcBLBxzQvH}6XZGV+=lC@0yV#3MVnmrPj;7d9V62(_4K|%+RgTkgZ@E#P_h2nFkEa6C2VMu7aQ67w(wMED*A4F=4GmGY1_2iO&#;k1AL#8L(<{Rl$_a+dlOZxfGo`$woZaP<T%tqYnI>_*U3d0>sqM(U@ugz2%E-*k_bqILo10IQpgK~250;YJ>;eicb_%J)<k3baJ$oT-R#GB<^$J4V5~aF714~Y?GXCsjj>G(CtI4AHu%IT=r+@t7b|$vyj6@f@}L<2C_W^cWt>yDjGGm)_aV_z9YBCRkedh8=^{D1``D-n!acjQmwwq3Wnluzs}Gl`f6hCmX<8LNeb~}-p~Uz@LK#kd(6MBL0rFtZv}GAAOs0p4{iI_IK|=>Im_Wc{#D<uQrbb`tRB%?Q8?cTETpghHoGjK#M|FKozGEm<=b5jpGYuUd_uADT{^j@onW&91(&7)+UsBli;k0fuf<{L(Ku$0$Wzm4fP-sU4nOIZSIUA}yPqDu5UpNw7AwkZo6?7iU?S_C^UtM+$GA$IL*jxe!(A;V5FA&`TzvrmbFW@XFbO+{TH7G1z$D=Ul9HUkYD+9z9!%EzNFDNQa@6q@CXTFCg0q0^3)?;*z>1C@2VL&0|P|&2Mof~VDsWp*TCQ}Gmq0Cs(FZjDZ{_Fq3gCCa0DZRGa=FwvX4?W0RadOPgWbNI-%ad#gd3gA4QO{u`SGS6hO)kXmHByL#@qeN~)D6wUpNeiVWAH(Z!9v4k>td?}vAv(&?ps2IGR-kRo+GAHy~}27vw6$Ol&mh2%H*q5(ZEXcxt(iyFT>wKP4*r<;LBxR$~fK<e%=9R*~Bc$t51@1Mo!WG7>mcLN{Sr}k{nSbEut1hKuG@ax#iZdOPec-&Az-~m{O$O`oeq4Xk8hhVpMZ%!n&dHY~Fd2I*u&^#*9f8YZ*EZQ|?$B=?p&6x8_2E@yxOZyGKmw*N|4|?X}a157Yia;TTcohYsaboI^yAq(kqq;l+W{Ai025s6xl)A2}^tfEGqtnos;5{tRiRcAe{pPRZ+)6P1@_vjIP5r-3cfwJ0~Qv0M7(@$lSm(A}<33X-E6y%_P-j?r+GZEIc#n=5Es(S5FqR$#v#cJY0_H~ZV)$Nd^^ECoU~6zcYe;ejdHc?<EyD6ac00xA@d2r5FHIfOi^y!cHnw3FS|ph4Uz-k~ZTrxxqXPd~+yG3gSZi3Fs^y=+!hwG9feI#-~;oV`r^T-cRH!H*kavRidC58%kV+y(=mdmOfOt0>Q^k<7E+NxHjNFy=|uVI{WA{8r$bn(2s91pC;T>7;(Det1kJNW>b%c6vPQ1!J6tB?e?nM^=^!^_m}o7<F98+}^fbQu&W(l{(o277kq@V9h@WHK1~%*|c)o&8-|PMaCQ*Sb6n`+!pqS6rZY(jH*77&2y5Y)R9}lv5g&lTh{QKjce7y9kGsBnmjc9nb7afMhoDSkbq#QX4O*|3Y2l2TU(u5Fjzq#d)t}8`$$7E$uY`}pr0$tM?PP4*16+kxfHAzdQ%|`24$NTHcAflpn8ljRTL6(#mZvZ`Z%{Zn!cgTT05zq5(Nv-xb_Hjew$!l09QP?{ftcYpy5K2k{y9&QD_$y_gpWzztM|JS^tw80~Kir{8CIXQP{AFB5;U^VO6|{nix#)B=euj5D0DGN~gyA=LyV!Zo2{-kcqW5Sy996p%b#nXwgy`!OLOlflY+H--5f2aAqAa&<RHfE<dZyY?e;$*6D}R=YfX|H~QK$e9P2oPSRX>vDqef6W-L;b?}gw^E@w+tk;puKi9M>b_-8~n}b288On|Y4ixP~vMi_!^L_8nqxY~|GqT-98#oY}93`x7mEJ=w%DFa}Z;tU8#yXNbo2)SVGM6Wogj5r{w86FXM40bd!-6g*lk>4~5-gU;a{_BbkBUrud0vAv%8O^eiXh8D`sdYb|2Cq5lV)38IcunI(p#d5RbqbOP;FW6JvC-MZxp5B7RbEf!8CX|Tc@~kQMPkDVR^URoM{3g4!6)2y(&?^7fMqZ?7d0m8VK%SR+1LIhhg%ZZT_VG;r`w0V^^0BLSF_RTvyF&RSAx(HU4c&q)$Y@<sW|kPl~E&gaZI};+$4L{ljm)b*>?yx7sI1`Zo8yc0hvwWeqkZUmx3c(7!@WBP~cnRoI+o#*<^!V(`L7-2i%=y{o`$VJm{5m0m+D@w~@^&jrcrm$Lk|CV#jv3715BnYZld&Yg>`Ln0ypfEdO4SzOVk1$nQafVE9gq^HHiyUF_NPHbs__A~6~Z6!vvP)!^T8Ydgp=B-ve7|eMgUHt@Wn9I_s>K+MULK!kNsP-5ecfJFFNsBMIhdlx$4`QQ{+vgB$v`U%U?MEAC&Lq={Uap^OpB%$S7js!A<m$#ENj9&DQw|E3uY4DiCJ1wenL!lu?Y`KFEHwtVY3n#?uUo3#+ME^vZJ^kH$gi;+7x<6!GJTx1Z;wO-Gds5E77r_;?jDWV=tB2TWS0R)hRlP7V{^AymO#My5r>}FsH99T1@jp-FAd<a18xArJ*a}lOa+)jA_1IwSe}Y=UAvJ0dU&~xqSyJ_aq_4Ql;=eCOr^c=;m|(@2c2$1yesSig#W%;dODnd{3!X%-I3Ot5$#SiJiAHEY-tI~9Nr!YDMH}i2E5C`3}E`eTv~CpK~~!GX?QM{(w7h^2#E()j!Pl&>ZXnbs9h2Af?FUgzIg@+N7iaG8lTpR?rJ5VzF~hh)W(Py8&+#XDDCQ-vKMRDVg2egE{xPFrofZoD9*D`K#Vn{iv2u1(xn9BDs}i^IZk==<m!WYY1HseUvUtL%UTVsRB<o@w4I?u(eNLCV;$Gnmywm_L=9@LhQtH@LJ;6v`380$+a$)!Q;oh6a7Q=-lLT%)1lWtQgr|2RZCx2omyk2Ijnj`8{49anBcj^B!P1Fa?A%l+IM)hwvwd4t*QPnNCmL59lmKqe5ldB!N11Uc*V5PI*U0^6#scGe3LTHz=vZlKX`x0r8K+guaV-H|0edFz5djo4|FYg;)_Y<9F-EmqZzHXIrzJFrb+giVdB!wB|F+p|SI86sHMz$V(4W^1j2Nv3D$OEoYRCL0xxu&Q9(#lD&+TszmC0p4=$)?b_V}S#n|katLf_!naTB<2BK6sUoKkV@vqf>wrsZDB$KHKoAGh;bDj2MVBwSAc+&R<%<$eW0>TO7fZbf9ib#l7DPw!{rQ`X|UyALqXyM`w=&d1ev)wd|ewll>hXMrZ|{(U4tT668AHMwVUr}Q0LKGa9m(nJ=@*A|BpJE;F~J9HklH@l>d{OFIM`;ni`R)wyLhLwJ(RR>pHlRDZ-JS3s!eAe)p5yXf#hKl+RhpffHZ7H>@Yyr)vZ)Voe8~rsyzE-Fs`gHjEfFoOKm2L)s+QRk?Od@u#;rfUEOt#Ri8q%#AbBYF6_N~nu$N$L57<xfPYTkoG^*6F`QWVg17|UpHwic<2lMS40JGKgPbEV5T0{m$shCut8mhn5KJ!|}`yqXYX>It@`%WIYB84etZ#O>2gd?sanNZdENmLl(vyP9WZgcEo&g>YSgkQ<@arV&Z9@Bx2Z9f?gpxSmGXt*1@SmS*vN<#Opd2ah`sbzNXFiH6%=f_umjbIz@JLbRNlkM8o!sE9gbZYp~xX~HymCXpfPp|Q!o(XXk*eLTA<eZ^=}l&_9!RU<@39tCc7-?s{9Z0w~VJ7pXVV@cX~mZvo6V>w(kuJ?BjS}rkyyZ6%eY<5$-%;_+B$FZE3m$sI>Sl))6^W=J#eQ@HP);m`nD~k8ciP$Jy3C<Iim2(M)SS_pkz{ag&-G8zJud*rLxzrOT-xZsH29oT=iMe}!W;Ipx^cfz8_+q!go~tI!xvZCekt=@@T`h%4u~l=fo0e*$gcvg0Mis<<;RCJoq#mS4xU{rHLo8q0ZCL`2-AG*(PpVJL<@-f#Jru)Vj7jTGu-M0`u|P&OB3|O!k@a~AqvbjZY%5k_oHrCjq3FzZZhot-{bQOKu*!f-ZJEEU+b=OTv(cEtu6_>R=#<YQMjvAByb~WKB1nA<R+vlgEb+_OJJd6LcKHPMjd<m!cVjqc(bR@V#hKbx3r!z%9#^uT5Qfe1HzqvPu+j)5-<ZeM^ugLZl1!pY@o_`!1^OIvus$;d_+GVw<Kw||V@5sKkfuF+)%pnUaf+_`D}omH+OH+s`_?3i-oWb{1YpN~rg<nz&x(=^n2Wj`{j<fX<_Wz1S@p3MBG76K_;E@rW-bSiSrdN2kVs?x{{XbMS~(%7E3N~o?^hc1SQtj`YD<$5ro4y_9gTCp9~5Agj$lB=dOYU^I5pnJJ}zc>ednn7`g=Croimnf3|ak`9FNi5oM>Drb@Zwhqk@N9^a@dvZu+l}k47}6Nq0lDo&0eN<b($`lB61Zek$_lAx=WriK31^X5^6JWFw=Hh-(@hD_oh)wjyM@I(9S)Wh->7)~T(AHjjymZOswjA!(qCQX3abQ;O9;&2(9}W;51yCofV*nd!Tgt6AcEf^)6%w%~XK6ee=-eU-w>@_E&h0)<UgVpRHAsU#~`WVcq)qp0e{ZapTU^etq(W2fh*ME^;lE4gJ<j+#td6y>8<=Dydk{&wwUU)}iiW!LbnId^u&`ROt;8{LXUY+Iy9lfbp^im&Fcos8fsQ+6rBHyUG9jkVMUt5tJzEXmajJVJj^3+pROoweAz#;jY#^TiU+j1$;-ZZYThdxz5d5GE<<$>q|BI2)fFPM89_Mvyg?9x-NEOc{?oSU@txt}Lara(|D%@@8^*`R%3%>O29k*ngx16e9<RHIKZFpEmO+8GW*NlxwSy^M9ssvjw1T1sB<;fBSHVOR6u550^SVI>NXFa-DW`vb$F!Rgn$m)m4%^XChQ1cxm!+^zkSQ%;}g*f&3ohG_GpUXrYX;Wzcfg<@`eQWZ0NmL)%O-Ev7?T#1-||$798r>$cWwit<Y-+?RmODU<P*@L4S%9r@nLiIJl*DWp`$@`mWS$1UdWSUxi~Eo&&aV~ypC54skQQ8JK9S(k!b0I5o-Meb<8k)@d{glSmC59_!<7K885fQG>y=dBEe)yc?1sAgs5ogK*xc|bIoGF>>d8ZPA)pasJADSp1~1r$ilC~$t#!h=T48g`2H2wi4vHFS}!(z`J^>$;#Hn$u-MDII^ZsOiZgMCAB{zZgL)!^^aw&*<@t?~Qqe0SZ!$8B>NM(1=>Ly)kwWgD#8yACUTkr#^|RFmu&z;!6XLCr^HhFD05H5@v3SyEryKA9;cGW@HWwXtUF|Czw%q4>dBDDR>jx<LxX}+1UQ?Giu%X9;-N`IgHh!cp;&Zv5yFgXP5BrN=I#q#pb0f$~bbUEvFIXr~(^h0zv6c3&+duE<YKeN1^-mEL1KFN<SxSCC@78EvRVfSge5Ku!dLZDe@4m4us{!23yf;Su`Z{yEW2@Dd}YW%~;IYPtqDbtCP7jNn$ggwlPVO&90maD|VYM`KDqYdjtxLk|{<5eSiDXRHOF5LH=!B?hOs#n+&L;BERz=B{dniBillVPjb3~m^{-X)YBBJ;-JP(-ano;vU!n|ud5#cm|mM*?2%~wwa4-AQYWZbc9e3Hj2FR<pbE179yM{lrQY~@*ey5u{0=2sPecV16XhQIzg?!+E1LVV0ROGnAbE#3>M>C4Xml{K+}t~PhfmuYcGrwbz6o#FVAIMdDMcYb>!LM<CYwdmz^d|UZ%y{!^zIMn)uFGbf=bF24nwb8?KEmqa0;1#+LP38^O#nhO7^iHEY^x2RqooZBm8jdPO$qJpXQW)v_GAG^_x9CIUqyHMA^DN9?D7omKL5t11q)gEc?mLeGri|^fU@VI{_SRULRa6INO707l}egeauS#&Vn!^xE?&j<6DX5uVawU8VFE?qZ`sJS`7kj_%!IAJJB_<@hHF<_g~nbbS#hh8{3DDf_p~VnZiYM@G<rd37H||YPA_($SWuL!~qv)Gfhsv$dzHU4=ohqNxEH(pzTtlc^h&YUFgq>lBkoAeg1s&Ojm`vgjq0D`1cg=paiGgPWt>cX<z4^+?98;_^ab!L1{+j_7ifD2{m{+Ocq|uz!w&$1V8qv^z&)<+)`Elc)0fo=B0L!*WQTNC=M0|Nh;4M!BRPOJ#)>r0XtHm9D5|8KFs+$pd4Lo6cqzcazB@Z$<O&X2&HKYTWSTNx(b6zU`kH4^?mE(xh6rNa`4nc!L>EMlYLP4LmpG^YSc|Px&VY#VUPMQx}KBTZD~ohuhMbUR+Gn7=yS4wdwbf5VmV!BDN?L$zNbgeLfdRd{n|sZ9mV$)mNzfuZ58dSEMUq)=hUs)JPUX)1!&3P%?U~Afz(%S<Qu^33Sh_Pj2N)r^?N#is3}|OIAL)Pwy-Jazr;xCa_1XQ462rJ(Ugv}q_<9$Bza%e+M!6((v~4Qc*v!|+HKv6QEqLc6qN70M!Dr}6(_VhKN+_ff6Vzkym0BLR5w>-5OxYugBAj5xV0n#MtuozdV)tOV=VwZN}|`NuxVp!_RGzqxoR<#64pvm%@$Tf_)?S2-8GAPG3{st*A3yqPuax>q_oKg)_o{LGZU0?P``n<ja5!d8oZ^4!iKi~HNj{TB8xPKSe|cu;yfZHhew17!S;Q6NsiWJ7p$6kyGauY22)5Kx4H+e@=PeR?AwUNptBx%t-Sud7?yk5+_7{t6+_UaP^(9S{K);K0631L2EaXKK-buvHuI<y*(J|lT$T(;bL&+D2Z%l-urR=E{bgg$Ei74=b6YU>)y(oN*5~2^%&Y%|x`*<71>aZcW}VF5K$Q{tge*iaR7>k{97$OXDp58cdlu3XWA?*ND)z$BsqNvyg^-a0&E@eat77YpN-4=a5(10Pj&jYCCF1!gS~dj6>9suTgI!vUks7SIyGHQ^9^K?JG?m@f%Q(L8Sv2n6lKTLHjP_WtK%qk_Z5iYP-fiWwnyTwQU=t;L{K)jqXt@S;Xx@req!pAl2UA~yrz!Fhgx`E=!ID$h)alk)`fZW$HEmalZDx{Qd4qX>_e;75-+I<)CC6yf#}Hir=LSL~fFdC7ZK|#FbnQuVPjcs1Y+6!s1t16^#*wHfv08;+US->0Z%?0*AEuk_>fGuy9PZefhX*eeT!9?t<`@Q`tC&pCt;YY`j$b%h80TE3!12!Iz>-R>Wb)C1A*~H7MW89iMFIA<wp+0_nkefp3S{ha-mF^pOzm<X@vq-?9x<NRKSp!F_A#T@QGa4qQ?f4>!ajn0yHfzWw|zJePlXz6ja*sLg8%8|X{f-ZH+S06UmkP62u2v()yg9`=8)pnc^KRrl(3Lc?>kLu>-b`I=2^nhY7J1ai;eXDrc9vBWqMl9`a2pus1B}|hI?Q6H{vai`xut^g1J5tsJE7qwf<PEi9Kvc#4>?&ei<&<`e&f+i)!E43YM{arRBOQe$E^-1o^oP$g#rR@)%vpBC9}^ZXRrIr5|QWWyeKNjs0Fw0ZE}px&67a2_uuRRyo7!amdzd*l4mzrK$Wy>WpAA49ewN{7STiO+M~iTxakhnG{_YEuuV>QlpiN>9?q6SYzXti(zHY`F&KcScwgr$f+R+9qR0R)L%+W@uHz56#hognSPAUNX{cLKfRWRvyI*Z#<SFlm1|`AvQNFxgP{!N1)fw4`pzs_<9GoVEEsY)k2dEZ-Y10eS>b!hL0|i|*XL~xMdetzwQA=Y_R1fBzCtM~7sxXTU;`d*&RA}zD+Y15xil7y3}{tnR^=ET&jIT+c|~gv+;MRABBBZI`uERb7^wc%$ojY0J{0TW=1lNFgH496nCRrHx)lQR-!qR0RM<VnSy%UnrBGY9o7_C|-RFgy#_scDxd<EQ#GcGFx_xt;v6^M6uTIH_%g24<SZt}oCOcZKOx>z9Im=yZEsh0;q{Sesgy35j?#pM^2MFzzix2vi^4d%uGHbRusUv+pdAA{K*&dIjElA#~dW09sf$5#RzFfh$v6Nca_^}#+t;5ry2$Niyn~v^j+urYSi;-^YY-L^Y+!flKYdgsDq(e8N*g!pYCQ?jKp$XA=x^X@EM)dT&CA!KrbOltTL*6Q}6pWeF(VQqBV1|nAm7F3~Xl8CST5OMmxMgvt{UgxcwA(L^thNyJZ6|HO&^R3Pcb&!EG+gI$?bs805}Tu0!Jqutm`7$Uchqs8FNo_SYidJt@AdJ>7N<Z!jNK8&8>$($Ec5t>=xI_n+~~{jy`KD{k9~@{b>`D;uA=cW?nLA?>wSa6k6oIf%WGR+*J`&MXx2rX)KY3J!|;8nQU8>XR~A;)n43ol=OT5aQ#IpCg{=+Es11Wet67n@qRgfHkNmFj=IX(RQrl*LS3|Gy;nuS#zU=s!KTcRSBA@m~G~CcCTI^a%L`SY7%Qmg=i{9y)bc4Hu-k6{RRsGoo_`CWzo2P2q)O=3t2b#Ak&SPFB)tIePvtCFCA}^^tr)xXqaB!IYSjJ!zj|b89Q5_qZHGN<EqyB}}UU6zeqf})=V;NPk(Qr70DqKp&1TstO*HIgACCV8D54{KXj<cObYi~ogA237IKP;G%@4Bc8bo4VHnpuJ0-KWqwLXSvKNCT0J5Xs}A^in?%FEqgkeNXK+xjwkx%})A-wzS$rbfkA^J~GeENXOLj4ds0DV*Q(POxSl?8!>LbgOhT7Ks^Y;YMvCit~Jhf0`AB9j^`X(h)nK;7-b9acMa=C&i10Bf>=~_2`gMOI+%GOix3pyi2b$6sCqG2xJllh#g5%<!#SaYtHSCc7AQVrZBezMtMBGYZ>f)vDgC%W7vWH@Ppc|5DBk~8OV<eZu_mTk>A_Xfdb5)&j|C$Pz6so0Cu6h$xNovu==ND|37(uxu~&DvRvqoq9eNOI!gMi-%Qx#rIg}#kM0TGIyV<(E#0uF95-P!Dooy_iGW{eLWpF*r!KbuDsI@M?TFh2AOAOuISdQzIbR_B<7Lc=EdB_1Rv0%{l%1sPBqame5ZP~^)0;z3F)f=y>dn>tVBh2(;(0JF&RC}8hy6h5>>!Av8rzy%YptBq4<uTib=B^yqUDv=Qxna?wJwysG=^JLd==8m9Bvx~_`J_#8VXP2r$(&raQ-?jMp;y<_HXjfl;zqL$oxufd|5Rfr3ZfF@Gd!n3Fuv2!vK$!_&0f`yQfpjpB&Bieh3gy5V!_$zHZ-_}OPp>|k)J^}T7O;b%NrkLyWCj%aWpS5)_xsj_raeZgCOWkrJ0sg{YX3iPuk!<sO9e(-J<PojRhk(WCTa>Z9z}lhTT%B5s(twHox6nIS(MSN;Bh>p?Y$uA+qsFga%kwKn^l6sz0Pb$bS6v_W)^&4zU+|rI|I4)GpN$*yFNnYHN2*YkyhZi2M!k%4x^U&}O6eb5d!Rymz9=$bF_&KYmgVWne|CFGvlD5V+3e>q9+Vl(n8MsOjBv0=HrYB>vWNR;AmpJiF^2>?>>{KVy|CQImu7OEanKU!c&=qjWg_P=$y&)E;Do*DdU4U)DK7zJ+3>?jB~R9sWAT^`?#MuosG_ddmkF>PQqcIDLvL8!d$V6|1+1pDQ_pwx`wXs*>J9-OQd_gVU)WQA4}}GB)J|U?dhQq1Mwrf#$r_BIC<#m6O?QdkrzdqSY4}#PD!ta5|qW-ui*DzqctqFLT-X3^WF{q|^KR*M}Q-5-dHjkgtRCPCu69Vc|RG4J@Atqm8xFN3y$eTz{e8ju=Hd^OgxqvS7(kEpZAmvVK>EtVc9pou4i3-*;1tEEjEJl1~7V7E_F5kH_XF<*rTTfCLVfpyg@1J$SkHiF9LO=>^@OZ#{oZAuTs~GAMd%W;bG|L3y_NLoVUQa5<VdRGO+;k1XLE8V!>`!=@Kw;FD>$ptr1O^i>OL7;>a#1sL?Be<0Qca(=TSN2zSrcglC@6yLt*N6LVVO&9mEa;LE7Y5ikw7@L>>#{C6&?gVJ(_}qPGg)2$njK4K=h3YYsZ8tDYMj<lvpj#^lrjxzwipI>AK4qMvRz3;w*q5<<!6F@&5mpT$50?i+B>~!k(@~F_UVs_~GgyzP|FeFVt$ZDH?tM8lvOR_s^tGT~&DtJY8O82k4=njcmf~PJ5uP&3zOza^<3W$E-ZreR(?uS-&`uo)dAb7g%cDgjVxAw18%1~ss|_jS#8pFLa3$7}aaB)Hj39njdpT;(W~z}tDsqka^IS=m!RP0TSwMjK<vM4ic7mHc!N|T6-@7A}HnO@WY{3Vf8k&#WBz53e=4t1QN-Ub<UMzBwTuCkA_8^K+03-ZAsrf{<OuQLjZ(7NLD2hCz^I{Q6@LN+!b%)3r)Hl+C>x6?ufJkhW4bETPLebD^6Vj1Xv;L{VT4O#_PLjgBxf)tUD6St$aMbRsrWn@(9ZD+9a!n6D#@9yuLC7s+WH}9Z+S&pY%T7)36sUbgCkSg+v7k+<Fcl|nWo>j&(m||(!~!sJ-=jakV<w5xj7VO~su5k^)gW`Y%De_f#O@OC89}+Z?H8b=4Bh)Az%p{siiH@x(1#aYXv8xX@R|<JN?4VHKa9cAy<NP;TdExSZ5W(IncN&Af~`p4VM?xR#Ncj`t&1#MmB=ttSUSVq{W6rT*i9~0ya}<X^I~<>e0s=t(4eH5<?`U}2cH~NN(BmZb@|$FmeO<Bd;&hTA4uELR2q<v4qwXRDRjm(w|5nFNiO!$@<S{qfV}jZ*b`5v?9IG^<a##3!7Ozn<Wg9Ygl0%YlC+cvO@<&+%#~azB*EAu>HA^Dogd4C@pNHzu|lO<eNoi_R;=YC3#H1Bt0k{{i2fSD+(gKjlJMuWx#zKCMD9{a@L;QaS4(JGg~dcPE%7mB5(v#s#wEp?XWULUvl^XJ4pdMHxldwLoJ}arQRP2@UQ!|9Q8?_2fo0)6;~+Sb@(}UV<a*<!Y{Wu^7!9FBp^Uyz+O}Scr!Se*(V81cwVm9L=7rhL<_V-aA8#92BGko($oEhs!0b+8&+0J5TU-_QqUuX~H`FlVKbdc56)&TjAC@H_M05-VT>LmnCzedZXbT?aS%A=B4u==?&ueI0o0PDgUMXdGGN7Y`T<|cB;89ScSJ4@!L@Hi`C=YBp+j;9I0vK4bfi-?lb~HAMUy9$8)p-|l9C-|yfua2XO7Nk5IGXiW$mbBGCnN$QacrobUThW)2`~skA1o=}Ko8XdisE@%X|n3U4HGd9!I5ikne4(<sH}J4n#ULc6X!85`P0HikMUE|XA9%5RwKNi^x^coYexk|;g<>#jPtv`9cBZs6He~KE6SYBbkGpYgn)}H906-rODKG=<i*?NEUf5+7PVj|SqWJjPMkA!)VLuVv1!*eA=0il3&ngI5p`h&Y84YmvtKl@@g>Hs=q?$ER>OULfa8?OX$mA$_Em3MFL^eJ)S;t2m@BF%bb>lP>>iGFSURElVAZu%j}LmQHF)(E&a7;hdNKd{xW^t1A#I9H*jNUlSgtN|M*fo(>wg=|ptG=#;c!cXHtM6*yhiF3R$XPOs;lsnH{@o$Ay>Id8=(>mLZH>8Nqkq%vXcU8Pa(BY1zb51Q(soB5&YIi7t+#4ve8T97Mvg~`>aY>VNkCw1qkbRQ(<-1)c9kf3umo0){C>Azz5!lyf9Y+hFWY2U|PI3oABUtee0lgFXyZHS+2=IuDY=%w=A^g6pYDVD5rms6s=m2!9e*+Rho<B;3{}AYzT?8z8@KnwlwLfE>k6&8D%i}^3vJJZHrATp4YXye}M@1MLH7@1&Fi$5uubOnBFzKSaf{>@g}ftR;;CGMa<LEJ1ow{Mg;N87!kzl_tVvRMvuTReC!oJm7<BtI9RB9{}BNtSCG8;Ws@RFRYZEAzu#*nK4Ux@-NEZM?rf2a97RNmb=sDZYeRcK_-`GlIS_jdi&=>}qw0z*cpx<w*(q=OGLP3LZ?u*3FAEA;EC4^Q<gpT=ONF-GQV~)d?z!fb*62-SILZ>IuYj!^kv{}>QT~{DpvK7<;ZwLtM2&D%q*v6nI=V>)zL7KYJ>poYlE5|RJK^mytywi3?{aZ<bW5!P_={>xsLEp?=G=S61>l}Lq!~YsRv3%Ev%+U^tQK8?)gr)r#K{`Ga&fn!+v@Fg2%22XOF%AyCV%i~<WM_{y;z!~+U>yNxK?<fH)Qk2y?ee$lurrLo>BV23Su9>V7X^*<px47dev38fPKOh9dT{bS55*PT{-SI46Un!3ClJ(g`;)UBgb-ZUL}FSog0h6Ox%v0J1E|@y={<;@W2+mN?rzy_?x|J%;)tNUHL_YZ)7Qr#$D0Z=v|NnrDrho<87E*Q1iH0sll@~=ya|tm?b@Q&xOL+a!StIu=>IrAJ=S0j;~b(V9SM(1{})L-YoweyAUSxN2n&lo{%*6)rvAzeect7<6EJNS83d_rFG4d;|kAmA;ieCfQwrNHtBF_7ooA2h+;rFLagt9-y40vQMVj{8@Q!>{U|$anpr^suLo2cbua+#I(fz7ZqXhql;GDq6YIL?(oq7SscJp^T%9$3k9$*S>&1>B$=<`@>T=L|rXP?^K?v_HSu-fX#M*`jhMCMxs)|Dc7dD4b0HcuWTV*l^J=gDe|94cHo4!|OVwbQ&T9VVH4CBMZ8dar7J~(gwMXf+p0ytk3EW+lqR%lggv&u09ZzhO@1)+uWC`p#PAv*QJk08%xpHFlj3;nu^2!K}7Tu2^nbHoA)xvILTL4KreX@_8U{Iws~PyZNIJigR0DqD!ovU6nGSz&Eo>h=VwYxx17GFY)lWC6dCofLVhu--raL~8fh0|EDnabQP$J{VTsC5IoZIw;po1x>Ir&fO9E89VUvFP!#pdzS7e1GFl3)edKDMXATKRBhfiHn${x#zy?CjXS-yztJ(bypkI`C}VtWT42Tb)ekG?YxE271dSa2pRsp8>)sW+2>nR^plqd|JivQ3tTkUBmT)EX67sQOzy8?YpLKu5NM-F6O;eAthJwMw20`q15=t+?iQHbWLgUKm#Xc7XL^2-vz+bTAzbJ=xW)7`Tv3<?h31uO?`4pS((4U3!4|I4voFHRY-IRQV^6Fo5vyUy;cFxFV%u>`~oE1l=J1H~<Mq|~D7G+*(@*sd;eizNv?=}{SEt29O9-fQIjPw?YWs#!BGcF7J7tqNs+LKTHn>^Ral9gs=`aN<T8gtf8(h9WC3T<e$K2#r|z{K<#b&e3~??gT47f{GA|M*)M#AF0BhA&x993K+!DN!A;7A#SxWypDR347N*fmM|6+TVdN2vZ>g7o}Zv05GwDUeTjU=JSW&{}ZcM^T4974oh+>)SDe!H7|9OuFQb#wGJ+MBbIzgtvyBe5^BonA}oj9EtL6S%M{%=s^6bKtNB;x&BoLp7qa#dPwOG+;*}sXSIVeR@tDoGYb_|2ERgkZH=*hdrh>)_4XfU%ei%0uRwRs74ade?a<!V?l~`U*eQCKlZ;1(2VCVCK)VdgT%)s@nb2EtA``;Qe9qV$zqPxhKZ?70;lWcF(;0KrIUFEGQ85qTwmTgmxxb?O%Q(F4RB~U;t<KW9z!3kh&R<~&dl0mbn8GmupiY-zGSD!55T8JQ|S_932#j`G+YT{ebP--cP^{Z1AW6k*J-ZRu$c-BRJuvE26EEUy9jXXIg1SEWC&eY{`Jqy6vMLVxGLv=m!40-WU(yr5t;U>bll0y`gVx}3dz#i>1krJ3k6vtE>iHCPwwGs3A>M#HzGsDPgj~QwLIc5sow0E{X0YqVrhvwk!YrZ7tc!$7^axAO8;Z{WrpEsIKMT)U$5l$5tB5fjobVHpu!=?qyhe;8|d~h`a_me%xFH7^u{4R^$Cvyie$9WXrQD7-D@XnbQ#HrIb^EKYk3aERGEI#FcfNK{=_AZseu%Zt<?Z`c5ez~^tgU;&sp^CQJ=sY!x0#(k4-eQT=*j^`@ZDzf%@S4*;(^b@U`af+ej@<"

_INK = "243447"
_BLUE = "1F4E79"
_BLUE_2 = "2F75B5"
_MUTED = "666666"
_CALLOUT = {
    "hard": ("FCE8E6", "A61B1B"),
    "suggest": ("E8F1FB", "1F5A93"),
    "pending": ("FFF2CC", "7F6000"),
}


def _manual_text() -> str:
    """解码内嵌手册全文；不访问文件系统。"""
    compressed = base64.b85decode(_MANUAL_TEXT_B85.encode("ascii"))
    return zlib.decompress(compressed).decode("utf-8-sig")


def _set_run_font(run, name="宋体", size=10.5, bold=False, color=_INK):
    """同时设置西文和东亚字体，降低不同 Word 版本的字体漂移。"""
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def _set_border(paragraph, color, side="left", size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def _keep(paragraph, keep_next=False):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    p_pr.append(OxmlElement("w:keepLines"))


def _add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    _set_run_font(run, size=9, color=_MUTED)


def _configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.35

    heading_specs = {
        "Heading 1": ("黑体", 16, _BLUE, 16, 8),
        "Heading 2": ("黑体", 13, _BLUE_2, 12, 5),
        "Heading 3": ("黑体", 11.5, _INK, 8, 4),
    }
    for style_name, (font, size, color, before, after) in heading_specs.items():
        style = document.styles[style_name]
        style.font.name = font
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def _classify(line):
    if re.match(r"^模块\d+\s", line):
        return "h1"
    if re.match(r"^[一二三四五六七八九十百]+、", line):
        return "h2"
    if re.match(r"^（[一二三四五六七八九十百]+）", line):
        return "h3"
    if line.startswith("【硬性规则"):
        return "hard"
    if line.startswith("【建议项"):
        return "suggest"
    if line.startswith("【待统一口径"):
        return "pending"
    return "body"


def _add_callout(document, text, kind):
    fill, foreground = _CALLOUT[kind]
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.18)
    paragraph.paragraph_format.right_indent = Cm(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    _set_shading(paragraph, fill)
    _set_border(paragraph, foreground)
    _keep(paragraph, keep_next=True)
    _set_run_font(
        paragraph.add_run(text), "黑体", 10.5, bold=True, color=foreground
    )


def _add_body(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.widow_control = True
    if text.startswith("□ "):
        paragraph.paragraph_format.left_indent = Cm(0.55)
        paragraph.paragraph_format.first_line_indent = Cm(-0.42)
        paragraph.paragraph_format.space_after = Pt(2.5)
    elif re.match(r"^\d+\.$", text):
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.space_after = Pt(0)
        _keep(paragraph, keep_next=True)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    _set_run_font(paragraph.add_run(text))


def generate_standard_manual() -> bytes:
    """生成完整标准手册并返回 DOCX 二进制字节。

    入参：无。
    返回：可直接作为 .docx 文件响应体的 bytes。
    """
    lines = [line.strip() for line in _manual_text().splitlines()]
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    _configure_styles(document)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(
        header.add_run("SAB星系列新闻稿统一标准手册"), size=8.5, color=_MUTED
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer)

    # 封面：标题与三项元数据均直接来自内嵌原文。
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(88)
    title.paragraph_format.space_after = Pt(12)
    _set_run_font(title.add_run(lines[0]), "黑体", 24, bold=True, color=_BLUE)
    for metadata in lines[1:4]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(5)
        label, value = metadata.split("：", 1)
        _set_run_font(paragraph.add_run(label + "："), "黑体", 10.5, bold=True)
        _set_run_font(paragraph.add_run(value))
    marker = document.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker.paragraph_format.space_before = Pt(28)
    _set_run_font(marker.add_run("内部工作参考文档"), "黑体", 11, bold=True, color=_MUTED)
    document.add_page_break()

    for line in lines[4:]:
        if not line:
            continue
        kind = _classify(line)
        if kind == "h1":
            document.add_page_break()
            paragraph = document.add_paragraph(line, style="Heading 1")
            _set_border(paragraph, _BLUE, side="bottom", size="10", space="5")
        elif kind == "h2":
            document.add_paragraph(line, style="Heading 2")
        elif kind == "h3":
            document.add_paragraph(line, style="Heading 3")
        elif kind in _CALLOUT:
            _add_callout(document, line, kind)
        elif line in {"使用说明", "手册执行结语"}:
            if line == "手册执行结语":
                document.add_page_break()
            paragraph = document.add_paragraph(line, style="Heading 1")
            _set_border(paragraph, _BLUE, side="bottom", size="10", space="5")
        else:
            _add_body(document, line)

    document.core_properties.title = "SAB星系列新闻稿统一标准手册"
    document.core_properties.subject = "SAB星系列新闻稿的选题、撰写、排版、审稿与发布"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


_VALID_COLUMNS = ("星动态", "星故事", "星分享", "星标杆", "星视频")


def _contains_any(text, words):
    return any(word in text for word in words)


def _title_and_body(article_text):
    lines = [line.strip() for line in article_text.splitlines() if line.strip()]
    return (lines[0] if lines else ""), "\n".join(lines[1:]), lines


def _has_internal_warning(text):
    """识别“内部使用、不得外传”的等义提示，不限定固定句式或标点。"""
    compact = re.sub(r"[\s，。；、:：/／\-—_（）()]+", "", text)
    internal_ok = "内部" in compact
    restriction_ok = _contains_any(
        compact,
        ("禁止对外转发", "请勿对外转发", "不得对外转发", "禁止外传", "请勿外传", "不得外传", "勿外传"),
    )
    return internal_ok and restriction_ok


def _first_content_paragraph(body):
    """取得正文首个实质段落，跳过内部提示语。"""
    for line in (part.strip() for part in body.splitlines()):
        if line and not _has_internal_warning(line):
            return line
    return ""


def _is_meeting_article(title):
    """以标题所表达的核心事件判断是否属于会议稿。"""
    return _contains_any(title, ("会议", "座谈会", "研讨会", "交流会", "工作会", "培训会", "总结会", "部署会"))


def _is_credit_line(line):
    """识别文末责任信息行，避免将供稿人等姓名纳入人物称谓检查。"""
    return bool(re.match(
        r"^\s*(?:文字|撰稿|供稿|摄影|摄像|配图|图片|编辑|排版|审核|审校)\s*[:：/／\-—]?\s*\S+",
        line,
    ))


def _result(rule_id, category, item, status, evidence, basis, tips, level="硬性规则"):
    """建立统一校验记录；status 仅使用四种固定值。"""
    return {
        "序号": rule_id,
        "规则类型": level,
        "检查类别": category,
        "检查项目": item,
        "结果": status,
        "问题证据": evidence,
        "标准依据": basis,
        "修改Tips": tips,
    }


def _general_checks(article_text, column_type, title, body, lines):
    """执行可由纯文本稳定判断的通用规则。"""
    results = []
    expected_tag = f"【{column_type}】"
    results.append(_result(
        "G01", "标题与栏目", "标题栏目标签与指定栏目一致",
        "通过" if title.startswith(expected_tag) else "不通过",
        title or "未发现标题",
        "模块2 一、标题栏目标签；模块7 六、标题与摘要检查",
        f"在标题开头使用“{expected_tag}”，并确保标签与正文核心内容一致。",
    ))

    banned = [word for word in ("成功举办", "顺利举办", "成功召开", "顺利召开", "圆满成功") if word in title]
    results.append(_result(
        "G02", "标题与栏目", "标题不使用自我评价式禁用措辞",
        "不通过" if banned else "通过",
        "、".join(banned) if banned else "标题未检出禁用措辞",
        "模块2 三、标题禁用措辞；模块7 二十七、一票退回红线",
        "删除禁用措辞；如需体现成效，改用经审核的具体成果或事实。",
    ))

    leadership_titles = (
        "董事长", "副董事长", "总裁", "副总裁", "总经理", "副总经理", "总裁助理",
        "党委书记", "党委副书记", "纪委书记", "工会主席", "总监", "领导",
    )
    leadership_pattern = re.compile(
        r"集团[^，。；\n]{0,20}(?:" + "|".join(map(re.escape, leadership_titles)) + r")"
    )
    group_hits = [line for line in lines if leadership_pattern.search(line)]
    results.append(_result(
        "G03", "组织称谓", "介绍公司领导时使用“总公司”而非“集团”",
        "不通过" if group_hits else "通过",
        "；".join(group_hits[:3]) if group_hits else "未检出使用“集团”介绍公司领导的情形",
        "模块4 三、“集团”统一改为“总公司”；模块7 八、姓名、职务与组织称谓检查",
        "介绍公司领导时将“集团”改为“总公司”；创业周年、发展历史等非领导身份介绍场景可以使用“集团”。",
    ))

    contact_patterns = {
        "手机号码": r"(?<!\d)1[3-9]\d{9}(?!\d)",
        "电子邮箱": r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
    }
    contact_hits = []
    for label, pattern in contact_patterns.items():
        if re.search(pattern, article_text):
            contact_hits.append(label)
    results.append(_result(
        "G04", "信息安全", "删除联系方式、个人隐私和其他可识别信息",
        "不通过" if contact_hits else "通过",
        "、".join(contact_hits) if contact_hits else "未检出手机号或电子邮箱格式",
        "模块4 四、客户和友商信息保护；模块7 十一、客户、友商和内部人员信息检查",
        "删除联系方式及可识别个人信息；确需公开时先完成授权与匿名处理。",
    ))

    # “客户”本身可能只是正常客情维护或生活化往来，不作为内部敏感关键词。
    sensitive_words = [word for word in ("经营", "财务", "人事", "资产", "生产技术", "订单", "市场策略") if word in article_text]
    warning = "内部通讯稿，禁止对外转发"
    warning_ok = _has_internal_warning(article_text)
    results.append(_result(
        "G06", "内部提示与发布范围", "重要会议或内部经营稿添加禁止外传提示",
        "需整改" if sensitive_words and not warning_ok else "通过",
        ("涉及：" + "、".join(sensitive_words)) if sensitive_words and not warning_ok else ("已包含内部使用、禁止外传的等义提示" if warning_ok else "未检出明显内部经营关键词"),
        "模块2 八、内部提示语；模块7 十五、内部提示与发布范围检查",
        f"确认稿件属性和发布范围；如属于重要会议或内部经营稿，在文章开头添加“{warning}”或意思一致的提示语。",
    ))

    time_ok = bool(re.search(r"(?:\d{4}年)?\d{1,2}月\d{1,2}日|近日|日前|当天", article_text))
    place_ok = bool(
        re.search(r"(?:在|于)[^，。；\n]{1,30}(?:召开|举行|开展|组织|举办|启动|完成)", article_text)
        or re.search(r"(?:莅临|前往|来到)[^，。；\n]{1,30}(?:开展|检查|指导|调研|考察|访问|交流)", article_text)
    )
    results.append(_result(
        "G09", "事实与新闻结构", "基本事实覆盖时间、地点、人物、事件等5W+1H要素",
        "需整改" if not (time_ok and place_ok) else "通过",
        f"时间线索：{'有' if time_ok else '未检出'}；地点线索：{'有' if place_ok else '未检出'}",
        "模块4 八、新闻时效与基本事实；模块7 七、事实与新闻结构检查",
        "核对Who、When、Where、What、Why、How，补齐影响理解的关键事实。",
    ))

    tail_text = "\n".join(lines[-20:])
    credit_aliases = {
        "文字": ("文字", "撰稿", "供稿"),
        "摄影": ("摄影", "摄像", "配图", "图片"),
        "编辑": ("编辑", "排版"),
        "审核": ("审核", "审校"),
    }
    credit_fields = [name for name, aliases in credit_aliases.items() if _contains_any(tail_text, aliases)]
    results.append(_result(
        "G10", "责任信息", "文末列明文字、摄影、编辑、审核责任人",
        "通过" if len(credit_fields) == 4 else "不通过",
        "已检出：" + ("、".join(credit_fields) if credit_fields else "无"),
        "模块3 七、文末供稿信息；模块7 二十五、责任信息与审核流程检查",
        "在文末补齐文字、摄影或配图、编辑或排版、审核等责任信息；允许使用含义对应的词汇，分隔形式不限。",
    ))

    body_lines = [line for line in lines[1:] if not _is_credit_line(line)]
    body_for_people = "\n".join(body_lines)
    person_end = r"(?=[，。；、]|表示|介绍|指出|认为|负责|参与|开展|完成|分享|汇报|提出|随后|为|$)"

    wrong_assistant_abbreviations = []
    assistant_pattern = re.compile(r"经理助理(?P<name>[\u4e00-\u9fff]{2,3})" + person_end)
    for match in assistant_pattern.finditer(body_for_people):
        name = match.group("name")
        wrong_abbreviation = name[0] + "助理"
        if wrong_abbreviation in body_for_people[match.end():]:
            wrong_assistant_abbreviations.append(f"{name}后文简称为“{wrong_abbreviation}”")
    results.append(_result(
        "G11", "人物称谓", "经理助理后续简称为“姓＋经理”",
        "需整改" if wrong_assistant_abbreviations else "通过",
        "；".join(wrong_assistant_abbreviations[:3]) if wrong_assistant_abbreviations else "未检出经理助理被简称为“姓＋助理”",
        "模块4 二、姓名和职务；补充执行口径：人物首次介绍与后续简称",
        "经理助理首次写明部门、完整职务和姓名，后续使用“姓＋经理”；例如“xx部经理助理路人甲”后文简称“路经理”。",
    ))

    specialist_pattern = re.compile(
        r"部[^，。；\n]{0,8}专员[\u4e00-\u9fff]{2,3}" + person_end
    )
    specialist_hits = [line for line in body_lines if specialist_pattern.search(line)]
    results.append(_result(
        "G12", "人物称谓", "无管理职务人物首次出现仅写部门和姓名",
        "需整改" if specialist_hits else "通过",
        "；".join(specialist_hits[:3]) if specialist_hits else "未检出“部门＋专员称谓＋姓名”的过度介绍",
        "模块4 二、姓名和职务；补充执行口径：普通人物首次介绍",
        "删除无管理职务人物的细化“专员”称谓，首次写部门和姓名即可；例如将“xx部xx专员路人甲”改为“xx部路人甲”。",
    ))
    return results


def _column_checks(article_text, column_type, title, body):
    """按栏目调用手册模块6的专属结构规则。"""
    if column_type == "星动态":
        results = []
        meeting_article = _is_meeting_article(title)
        first_paragraph = _first_content_paragraph(body)
        time_ok = bool(re.search(r"(?:\d{4}年)?\d{1,2}月\d{1,2}日|近日|日前|当天", first_paragraph))

        event_words = ("召开", "举行") if meeting_article else ("莅临", "指导", "检查", "调研", "考察", "开展", "启动", "完成")
        event_ok = _contains_any(first_paragraph, event_words)
        results.append(_result(
            "C01", f"{column_type}专属结构", "首段交代事件时间和核心事项",
            "通过" if time_ok and event_ok else "需整改",
            f"首段时间线索：{'有' if time_ok else '无'}；核心事项线索：{'有' if event_ok else '无'}",
            "模块6 二、（三）摘要",
            "在首段补充事件时间和核心事项；地点、人物及议题按事件实际情况交代。",
            level="栏目规则",
        ))

        if meeting_article:
            checks = {
                "时间": time_ok,
                "地点": bool(re.search(r"(?:在|于)[^，。；\n]{1,30}(?:召开|举行)", first_paragraph)),
                "会议名称": _contains_any(first_paragraph, ("会议", "座谈会", "研讨会", "交流会", "工作会", "培训会", "总结会", "部署会")),
                "主持人": "主持" in first_paragraph,
                "参会范围": _contains_any(first_paragraph, ("参加", "出席", "参会")),
            }
            missing = [name for name, ok in checks.items() if not ok]
            results.append(_result(
                "C02", f"{column_type}专属结构", "会议稿首段包含会议基本信息",
                "需整改" if missing else "通过",
                "缺少：" + "、".join(missing) if missing else "会议首段基本要素齐全",
                "模块6 二、（四）正文首段",
                "仅会议稿需要在首段交代时间、地点、会议名称、主持人及必要参会范围；补齐当前缺失要素。",
                level="栏目规则",
            ))
        else:
            results.append(_result(
                "C02", f"{column_type}专属结构", "会议稿首段包含会议基本信息",
                "通过", "标题核心事件不属于会议稿，本项不适用",
                "模块6 二、（四）正文首段",
                "无需补充会议名称、主持人或参会范围；按实际事件写清时间、地点、人物和事项即可。",
                level="栏目规则",
            ))

        followup_ok = _contains_any(article_text, ("部署", "要求", "下一步", "后续", "行动", "任务", "建议", "指导"))
        results.append(_result(
            "C03", f"{column_type}专属结构", "正文提炼主要内容和后续行动",
            "通过" if followup_ok else "需整改",
            "已检出要求或行动线索" if followup_ok else "未检出要求或后续行动线索",
            "模块6 二、（五）正文主体、（六）结尾",
            "提炼事件核心内容、相关要求及后续行动，避免机械罗列流程。",
            level="栏目规则",
        ))
        return results

    rules = {
        "星故事": [
            ("C01", "叙事包含真实困难或阻力", ("困难", "挑战", "压力", "瓶颈", "阻力"), None, "模块6 五、（一）选题标准、（二）核心叙事链", "用具体事实说明任务中的困难和矛盾，不夸大冲突。"),
            ("C02", "叙事包含选择、行动和转折", ("选择", "决定", "行动", "攻坚", "突破", "转折"), None, "模块6 五、（二）核心叙事链", "写清关键时刻的判断、连续行动及改变局面的转折。"),
            ("C03", "结果真实且价值升华由事实自然形成", ("成果", "结果", "收获", "成长", "阶段"), None, "模块6 五、（一）选题标准、（三）写作要求", "说明真实、合规的阶段结果，再从事实自然总结团队或SAB人精神价值。"),
        ],
        "星分享": [
            ("C01", "案例交代背景、问题或任务", ("背景", "问题", "任务", "挑战", "目标"), None, "模块6 四、（四）STAR模型、（十）标准成稿结构", "补充案例情境、具体问题、目标和限制条件。"),
            ("C02", "案例写清关键行动及背后判断", ("行动", "措施", "做法", "步骤", "选择", "方法"), None, "模块6 四、（四）STAR模型、（五）5Why分析", "说明采取了什么行动、为什么选择该方案及方案有效原因。"),
            ("C03", "案例包含结果、复盘与复制边界", ("结果", "成果", "复盘", "复制", "适用", "风险"), None, "模块6 四、（七）至（十）", "补充合规结果证据、成功因素、适用条件、风险提示和不可机械照搬的边界。"),
        ],
        "星标杆": [
            ("C01", "人物或集体具有真实可核验事迹", ("事迹", "行动", "负责", "完成", "攻坚", "贡献"), None, "模块6 三、（一）人物选择", "用可核验的具体行动和故事说明代表性，不只罗列身份或形容词。"),
            ("C02", "稿件体现榜样和示范价值", ("榜样", "示范", "带动", "影响", "学习", "代表"), None, "模块1 五、星标杆；模块6 三、（一）人物选择", "说明人物或集体为什么值得学习，以及代表的岗位、团队或精神。"),
            ("C03", "人物信息、单位和岗位准确", ("公司", "单位", "部门", "岗位", "职务"), None, "模块6 三、（二）多人系列海报；模块4 二、姓名和职务", "核对姓名、单位、岗位与最新正式职务，不自行包装岗位级别。"),
        ],
        "星视频": [
            ("C01", "发布文字包含主题标题和背景说明", ("背景", "主题", "为", "近日"), None, "模块6 六、星视频最低发布结构", "补充栏目标签、主题标题和一句背景说明。"),
            ("C02", "包含必要人物、地点和日期说明", ("月", "日", "在", "于"), None, "模块6 六、星视频最低发布结构", "按内容需要补充人物、地点和日期说明。"),
            ("C03", "列明摄像、剪辑和审核信息", ("摄像：", "摄影："), ("剪辑：", "审核："), "模块6 六、星视频最低发布结构", "在发布配文末尾补齐文字、摄影或摄像、剪辑、审核信息。"),
        ],
    }
    results = []
    for rule_id, item, required_a, required_b, basis, tips in rules[column_type]:
        a_ok = _contains_any(article_text, required_a)
        if required_b is None:
            ok = a_ok
            evidence = "已检出相关结构线索" if ok else "未检出相关结构线索"
        elif isinstance(required_b, str):
            b_ok = required_b in article_text
            ok = a_ok and b_ok
            evidence = f"要素A：{'有' if a_ok else '无'}；{required_b}：{'有' if b_ok else '无'}"
        else:
            b_ok = all(word in article_text for word in required_b)
            ok = a_ok and b_ok
            evidence = f"要素组A：{'有' if a_ok else '无'}；要素组B：{'齐全' if b_ok else '不齐全'}"
        results.append(_result(
            rule_id, f"{column_type}专属结构", item,
            "通过" if ok else "需整改", evidence, basis, tips,
            level="栏目规则",
        ))
    return results


def _review_docx(article_text, column_type, results, counts):
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.right_margin = Cm(2.54)
    section.bottom_margin = section.left_margin = Cm(2.54)
    _configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run("稿件写作整改提示"), "黑体", 20, bold=True, color=_BLUE)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    _set_run_font(subtitle.add_run(f"栏目：{column_type}"), "宋体", 10.5, color=_MUTED)

    summary = document.add_paragraph()
    summary.paragraph_format.space_after = Pt(12)
    _set_shading(summary, "E8F1FB")
    _set_border(summary, _BLUE_2)
    _set_run_font(summary.add_run(
        f"校验概览：明确问题 {counts['不通过'] + counts['需整改']} 项；"
        f"通过自动校验 {counts['通过']} 项。"
    ), "黑体", 10.5, bold=True, color=_BLUE)

    if not results:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run("未发现可由当前文本规则识别的问题。"), "宋体", 10.5)
    for index, result in enumerate(results, 1):
        heading = document.add_paragraph(style="Heading 2")
        heading.add_run(f"{index}. [{result['结果']}] {result['检查项目']}")
        for label, value in (
            ("问题", result["问题证据"]),
            ("标准依据", result["标准依据"]),
            ("修改Tips", result["修改Tips"]),
        ):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.35)
            _set_run_font(paragraph.add_run(label + "："), "黑体", 10.5, bold=True)
            _set_run_font(paragraph.add_run(value))

    appendix = document.add_paragraph("原稿", style="Heading 1")
    _set_border(appendix, _BLUE, side="bottom", size="10", space="5")
    for line in article_text.splitlines() or [""]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        _set_run_font(paragraph.add_run(line))

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _review_xlsx(column_type, results):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "校验清单"
    headers = ["序号", "规则类型", "检查类别", "检查项目", "结果", "问题证据", "标准依据", "修改Tips"]
    sheet.append(["SAB星系列新闻稿结构化校验清单"])
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
    sheet["A1"].font = Font(name="微软雅黑", size=16, bold=True, color="FFFFFF")
    sheet["A1"].fill = PatternFill("solid", fgColor=_BLUE)
    sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[1].height = 28
    sheet.append([f"栏目：{column_type}"])
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))
    sheet["A2"].font = Font(name="宋体", size=10, color=_MUTED)
    sheet.append(headers)
    for result in results:
        sheet.append([result[header] for header in headers])

    header_fill = PatternFill("solid", fgColor=_BLUE_2)
    thin = Side(style="thin", color="D9E2F3")
    status_fills = {
        "通过": PatternFill("solid", fgColor="E2F0D9"),
        "不通过": PatternFill("solid", fgColor="FCE8E6"),
        "需整改": PatternFill("solid", fgColor="FFF2CC"),
    }
    for cell in sheet[3]:
        cell.font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in sheet.iter_rows(min_row=4, max_row=sheet.max_row, min_col=1, max_col=len(headers)):
        for cell in row:
            cell.font = Font(name="宋体", size=10)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
        row[4].fill = status_fills.get(row[4].value, PatternFill())
        row[4].font = Font(name="微软雅黑", size=10, bold=True)
        row[4].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    widths = [9, 12, 18, 32, 14, 36, 42, 42]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:H{sheet.max_row}"
    sheet.sheet_view.showGridLines = False
    sheet.print_title_rows = "1:3"
    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.page_margins.left = sheet.page_margins.right = 0.25
    sheet.page_margins.top = sheet.page_margins.bottom = 0.5

    stream = BytesIO()
    workbook.save(stream)
    return stream.getvalue()


def review_news_article(article_text: str, column_type: str) -> dict:
    """按手册审核稿件，仅返回明确问题的汇总与结构化记录。

    返回结构：
    {
        "summary": {各状态数量},
        "results": [结构化校验记录]
    }

    本函数不执行 DOCX/XLSX 渲染，也不生成任何附件。
    """
    if not isinstance(article_text, str) or not article_text.strip():
        raise ValueError("article_text 必须为非空字符串")
    if column_type not in _VALID_COLUMNS:
        raise ValueError("column_type 必须为：星动态、星故事、星分享、星标杆、星视频")

    title, body, lines = _title_and_body(article_text)
    all_results = _general_checks(article_text, column_type, title, body, lines)
    all_results.extend(_column_checks(article_text, column_type, title, body))
    # 对外成果只呈现明确问题，不保留不确定状态或相关汇总字段。
    results = [result for result in all_results if result["结果"] in {"不通过", "需整改"}]
    summary = {
        "明确问题": len(results),
        "不通过": sum(r["结果"] == "不通过" for r in results),
        "需整改": sum(r["结果"] == "需整改" for r in results),
        "通过": sum(r["结果"] == "通过" for r in all_results),
    }
    return {
        "summary": summary,
        "results": results,
    }


__all__ = ["generate_standard_manual", "review_news_article"]
